"""Document loaders for PDF, DOCX, and image files."""

from __future__ import annotations

from pathlib import Path

from pydantic import BaseModel, Field

from utils.logging import get_logger

logger = get_logger(__name__)

# All file extensions supported by the ingestion pipeline
SUPPORTED_EXTENSIONS: set[str] = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".markdown",
    ".png",
    ".jpg",
    ".jpeg",
    ".tiff",
    ".bmp",
}

_IMAGE_EXTENSIONS: set[str] = {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}


class LoadedDocument(BaseModel):
    """Represents a loaded document segment ready for further processing.

    Attributes:
        text: Extracted text content from the document segment.
        page_number: Page number (0-indexed). 0 for formats without pages.
        source_file: Original file path.
        file_type: Type of the source file (pdf/docx/image).
        metadata: Additional metadata from the loader.
    """

    text: str
    page_number: int = 0
    source_file: str
    file_type: str
    metadata: dict = Field(default_factory=dict)


def load_pdf(file_path: str | Path) -> list[LoadedDocument]:
    """Load a PDF file and extract text page by page using PyMuPDF.

    Args:
        file_path: Path to the PDF file.

    Returns:
        List of LoadedDocument instances, one per page.

    Raises:
        FileNotFoundError: If the file does not exist.
        RuntimeError: If PDF parsing fails.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    documents: list[LoadedDocument] = []

    try:
        import fitz  # PyMuPDF

        with fitz.open(str(path)) as doc:
            logger.info("loading_pdf", file=str(path), pages=len(doc))
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                documents.append(
                    LoadedDocument(
                        text=text.strip(),
                        page_number=page_num,
                        source_file=str(path),
                        file_type="pdf",
                        metadata={"total_pages": len(doc)},
                    )
                )
    except Exception as exc:
        logger.error("pdf_load_failed", file=str(path), error=str(exc))
        raise RuntimeError(f"Failed to load PDF: {path}") from exc

    return documents


def load_docx(file_path: str | Path) -> list[LoadedDocument]:
    """Load a DOCX file and extract text from all paragraphs.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        List containing a single LoadedDocument with all text.

    Raises:
        FileNotFoundError: If the file does not exist.
        RuntimeError: If DOCX parsing fails.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n".join(paragraphs)
        logger.info("loading_docx", file=str(path), paragraphs=len(paragraphs))

        return [
            LoadedDocument(
                text=full_text,
                page_number=0,
                source_file=str(path),
                file_type="docx",
                metadata={"paragraph_count": len(paragraphs)},
            )
        ]
    except Exception as exc:
        logger.error("docx_load_failed", file=str(path), error=str(exc))
        raise RuntimeError(f"Failed to load DOCX: {path}") from exc


def load_image(file_path: str | Path) -> list[LoadedDocument]:
    """Load an image file placeholder (OCR will handle text extraction).

    Args:
        file_path: Path to the image file.

    Returns:
        List containing a single LoadedDocument with empty text and OCR flag.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Image file not found: {path}")

    logger.info("loading_image", file=str(path), note="OCR needed for text extraction")

    return [
        LoadedDocument(
            text="",
            page_number=0,
            source_file=str(path),
            file_type="image",
            metadata={"ocr_needed": True},
        )
    ]


def load_text(file_path: str | Path) -> list[LoadedDocument]:
    """Load a plain text file.

    Args:
        file_path: Path to the text file.

    Returns:
        List containing a single LoadedDocument with all text.

    Raises:
        FileNotFoundError: If the file does not exist.
        RuntimeError: If text reading fails.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Text file not found: {path}")

    try:
        text = path.read_text(encoding="utf-8")
        logger.info("loading_text", file=str(path), chars=len(text))

        return [
            LoadedDocument(
                text=text,
                page_number=0,
                source_file=str(path),
                file_type="txt",
                metadata={"encoding": "utf-8"},
            )
        ]
    except Exception as exc:
        logger.error("text_load_failed", file=str(path), error=str(exc))
        raise RuntimeError(f"Failed to load text file: {path}") from exc


def load_document(file_path: str | Path) -> list[LoadedDocument]:
    """Factory function to load a document based on its file extension.

    Detects the file type by extension and dispatches to the appropriate loader.

    Args:
        file_path: Path to the document file.

    Returns:
        List of LoadedDocument instances.

    Raises:
        ValueError: If the file extension is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file extension: '{ext}'. "
            f"Supported extensions: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info("load_document_dispatching", file=str(path), extension=ext)

    if ext == ".pdf":
        return load_pdf(path)
    elif ext in {".docx", ".doc"}:
        return load_docx(path)
    elif ext in {".txt", ".md", ".markdown"}:
        # Markdown is plain text for retrieval purposes — the `[N]` citation
        # contract works on raw text, and chunking is markup-agnostic.
        return load_text(path)
    elif ext in _IMAGE_EXTENSIONS:
        return load_image(path)
    else:
        raise ValueError(f"Unsupported file extension: '{ext}'")
